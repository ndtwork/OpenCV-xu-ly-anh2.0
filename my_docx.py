from docx import Document
from docx.shared import Pt

def create_software_spec_doc():
    document = Document()
    document.add_heading('Tài liệu đặc tả phần mềm', 0)

    # Giới thiệu
    document.add_heading('Giới thiệu', level=1)
    document.add_paragraph(
        "Tài liệu này cung cấp đặc tả chi tiết về Ứng dụng Lọc Ảnh. "
        "Ứng dụng cho phép người dùng áp dụng các bộ lọc khác nhau lên ảnh sử dụng OpenCV và PyQt5 cho giao diện đồ họa (GUI)."
    )

    # Tính năng
    document.add_heading('Tính năng', level=1)
    document.add_paragraph(
        "- Nhập và hiển thị ảnh\n"
        "- Áp dụng các loại bộ lọc khác nhau:\n"
        "  - Bộ lọc thông thấp\n"
        "  - Bộ lọc thông cao\n"
        "  - Bộ lọc đạo hàm\n"
        "  - Bộ lọc trung vị\n"
        "  - Bộ lọc cực tiểu\n"
        "  - Bộ lọc cực đại\n"
        "  - Bộ lọc đa tùy chỉnh\n"
        "- Lưu ảnh đã lọc\n"
        "- Xuất ảnh ra DOCX"
    )

    # Cài đặt
    document.add_heading('Cài đặt', level=1)
    document.add_paragraph(
        "### Yêu cầu\n"
        "- Python 3.12\n"
        "- Miniconda hoặc Anaconda\n\n"
        "### Thiết lập\n"
        "1. Clone repository:\n"
        "    ```sh\n"
        "    git clone <repository-url>\n"
        "    cd <repository-directory>\n"
        "    ```\n"
        "2. Tạo và kích hoạt môi trường conda:\n"
        "    ```sh\n"
        "    conda env create -f environment.yml\n"
        "    conda activate ImageFilteringApp\n"
        "    ```\n"
        "3. Cài đặt các gói phụ thuộc bổ sung nếu cần:\n"
        "    ```sh\n"
        "    pip install -r requirements.txt\n"
        "    ```"
    )

    # Sử dụng
    document.add_heading('Sử dụng', level=1)
    document.add_paragraph(
        "1. Chạy ứng dụng:\n"
        "    ```sh\n"
        "    python main.py\n"
        "    ```\n"
        "2. Sử dụng GUI để nhập ảnh, áp dụng bộ lọc và lưu kết quả."
    )

    # Cấu trúc dự án
    document.add_heading('Cấu trúc dự án', level=1)
    document.add_paragraph(
        "- `main.py`: Điểm vào của ứng dụng.\n"
        "- `nonlinear_filters.py`: Chứa các hàm áp dụng bộ lọc phi tuyến.\n"
        "- `environment.yml`: File cấu hình môi trường conda.\n"
        "- `README.md`: Tài liệu dự án."
    )

    # Đóng góp
    document.add_heading('Đóng góp', level=1)
    document.add_paragraph(
        "1. Fork repository.\n"
        "2. Tạo nhánh mới (`git checkout -b feature-branch`).\n"
        "3. Commit thay đổi của bạn (`git commit -am 'Add new feature'`).\n"
        "4. Push lên nhánh (`git push origin feature-branch`).\n"
        "5. Tạo Pull Request mới."
    )

    # Giấy phép
    document.add_heading('Giấy phép', level=1)
    document.add_paragraph("Dự án này được cấp phép theo giấy phép MIT.")

    # Lưu tài liệu
    document.save('Tai_lieu_dac_ta_phan_mem.docx')

if __name__ == "__main__":
    create_software_spec_doc()